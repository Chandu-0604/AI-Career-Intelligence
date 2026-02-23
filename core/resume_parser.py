import pdfplumber
import docx2txt
import os
from core.text_preprocessor import clean_resume_text
from docx import Document

# ---------------- PDF EXTRACTION ----------------
def extract_text_from_pdf(pdf_path):
    text = ""

    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

    except Exception as e:
        print("PDF extraction failed:", e)

    return text


# ---------------- DOCX EXTRACTION (ADVANCED) ----------------
def extract_text_from_docx(docx_path):

    combined_text = ""

    try:
        # Method 1: docx2txt (good for headers, footers, textboxes)
        text1 = docx2txt.process(docx_path)

        if text1:
            combined_text += text1 + "\n"

        # Method 2: python-docx (good for paragraphs & tables)
        doc = Document(docx_path)

        for para in doc.paragraphs:
            if para.text.strip():
                combined_text += para.text.strip() + "\n"

        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    combined_text += row_text + "\n"

        return combined_text

    except Exception as e:
        print("DOCX extraction failed:", e)
        return ""


# ---------------- MAIN ROUTER ----------------
def extract_resume_text(file_path):

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)

    elif ext == ".docx":
        text = extract_text_from_docx(file_path)

    else:
        return None

    # ---------- CLEANUP (KEEP NEWLINES) ----------
    if not text:
        return ""

    # keep RAW text for skill matching
    raw_text = text

    # cleaned version only for AI
    cleaned_text = clean_resume_text(text)

    return raw_text